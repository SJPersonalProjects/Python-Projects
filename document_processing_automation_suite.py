# Document Processing Automation Suite.

import os
import shutil
import PyPDF2
import docx

# Ask the user for the folder.
folder = input("Enter the folder path: ")

# Create the backup folder.
backup_folder = os.path.join(folder, "Backup")
os.makedirs(backup_folder, exist_ok=True)

# Create summary document.
summary = docx.Document()
summary.add_heading("Document Processing Report", level=0)

processed_files = 0

# Scan every file.
for filename in os.listdir(folder):

    file_path = os.path.join(folder, filename)

    # Skip folders.
    if os.path.isdir(file_path):
        continue

    # ----------------------- PDF -----------------------
    if filename.lower().endswith(".pdf"):

        summary.add_heading(filename, level=1)

        with open(file_path, "rb") as pdf_file:
            reader = PyPDF2.PdfReader(pdf_file)

            text = ""

            for page in reader.pages:
                page_text = page.extract_text()

                if page_text:
                    text += page_text

        
        if text:
            summary.add_paragraph(text[:500])
        else:
            summary.add_paragraph("No readable text found.")

        shutil.copy(file_path, backup_folder)
        processed_files += 1

    
    # -------------------- DOCX ------------------------
    elif filename.lower().endswith(".docx"):

        summary.add_heading(filename, level=1)

        document = docx.Document(file_path)

        text = ""

        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"

        if text.strip():
            summary.add_paragraph(text[:500])
        else:
            summary.add_paragraph("Document is empty.")

        shutil.copy(file_path, backup_folder)
        processed_files += 1

# Save report.
report_path = os.path.join(folder, "Summary_Report.docx")
summary.save(report_path)

print(f"\nProcessed {processed_files} documents.")
print("Summary report created successfully!")
print("Backup folder created.")
