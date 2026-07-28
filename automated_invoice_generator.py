# Automated Invoice Generator.

import docx

# Customer data.
customers = [
    {
        "name": "Ali",
        "address": "Hyderabad",
        "items": [
            ("Laptop", 1, 800),
            ("Mouse", 2, 20)
        ]
    },
    {
        "name": "Sara",
        "address": "Karachi",
        "items": [
            ("Keyboard", 1, 45),
            ("Monitor", 1, 180),
            ("Mouse Pad", 2, 10)
        ]
    }
]

invoice_number = 1

# Generate an invoice for each customer.
for customer in customers:
    
    document = docx.Document()

    # Invoice title.
    document.add_heading(f"Invoice #{invoice_number}", level=0)

    # Customer information.
    document.add_heading("Customer Information", level=1)
    document.add_paragraph(f"Name: {customer['name']}")
    document.add_paragraph(f"Address: {customer['address']}")

    # Items heading.
    document.add_heading("Purchased Items", level=1)

    # Create table.
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    header = table.rows[0].cells
    header[0].text = "Item"
    header[1].text = "Quantity"
    header[2].text = "Price"
    header[3].text = "Total"

    grand_total = 0

    # Add purchased items.
    for item_name, quantity, price in customer["items"]:

        total = quantity * price
        grand_total += total

        row = table.add_row().cells
        row[0].text = item_name
        row[1].text = str(quantity)
        row[2].text = f"${price}"
        row[3].text = F"${total}"

    # Grand total.
    document.add_heading("Invoice Total", level=1)
    document.add_paragraph(f"Grand Total: ${grand_total}")

    # Save invoice.
    filename = f"Invoice_{invoice_number}_{customer['name']}.docx"
    document.save(filename)

    print(f"Created {filename}")

    invoice_number += 1

print("\nAll invoices generated successfully!!")