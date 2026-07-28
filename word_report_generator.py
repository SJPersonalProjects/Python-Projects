# Word Report Generator.
import docx

# Create a new Word document.
document = docx.Document()

# Get the user input.
title = input("Enter the report title: ")
name = input("Enter your name: ")
department = input("Enter department: ")
summary = input("Enter report summary: ")

# Add title.
document.add_heading(title, level=0)

# Add sections.
document.add_heading("Prepared By", level=1)
document.add_paragraph(name)

document.add_heading("Department", level=1)
document.add_paragraph(department)

document.add_heading("Summary", level=1)
document.add_paragraph(summary)

# Save the document.
document.save("report.docx")

print("\nReport created successfully!")
print("Saved as 'report.docx'")